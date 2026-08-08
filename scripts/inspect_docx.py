#!/usr/bin/env python3
"""Emit a compact, ordered inventory of DOCX text and layout features."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


def paragraph_record(paragraph):
    return {
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment),
        "text": paragraph.text,
    }


def table_record(table):
    return {
        "rows": len(table.rows),
        "cols": len(table.columns),
        "style": table.style.name if table.style else None,
        "cells": [[cell.text for cell in row.cells] for row in table.rows],
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    document = Document(args.docx)
    sections = []
    for section in document.sections:
        sections.append(
            {
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "right_margin": section.right_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "header_distance": section.header_distance,
                "footer_distance": section.footer_distance,
                "header": [paragraph_record(p) for p in section.header.paragraphs],
                "footer": [paragraph_record(p) for p in section.footer.paragraphs],
            }
        )

    payload = {
        "path": str(args.docx),
        "sections": sections,
        "inline_shapes": len(document.inline_shapes),
        "paragraphs": [paragraph_record(p) for p in document.paragraphs],
        "tables": [table_record(t) for t in document.tables],
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(payload, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
