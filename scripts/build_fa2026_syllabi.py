#!/usr/bin/env python3
"""Build synchronized Fall 2026 BUS210 DOCX and Canvas-ready HTML syllabi."""

from __future__ import annotations

import argparse
import html
import re
from dataclasses import dataclass
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "deliverables" / "fall-2026"
PHOTO = ROOT / "assets" / "img" / "table_Finance.jpg"

NAVY = "0A2540"
BLUE = "1D4ED8"
TEAL = "0F766E"
GOLD = "B7791F"
TERRA = "B4533C"
INK = "172033"
SLATE = "475569"
LINE = "CBD5E1"
PALE_BLUE = "EFF6FF"
PALE_TEAL = "F0FDFA"
PALE_GOLD = "FFFBEB"
PALE_TERRA = "FFF1ED"
PALE_GRAY = "F8FAFC"
WHITE = "FFFFFF"

ACADEMIC_CALENDAR_URL = (
    "https://www.endicott.edu/academics/academic-resources-support/academic-calendar/undergraduate"
)
ACCESSIBILITY_URL = (
    "https://www.endicott.edu/academics/academic-resources-support/division-of-academic-success/"
    "center-for-accessibility-services"
)
ACADEMIC_SUCCESS_URL = (
    "https://www.endicott.edu/academics/academic-resources-support/division-of-academic-success"
)
AI_GUIDANCE_URL = (
    "https://www.endicott.edu/academics/provost/"
    "guidance-for-using-artificial-intelligence-at-endicott-college"
)
TITLE_IX_URL = "https://www.endicott.edu/title-ix"
OFFICE_HOURS_URL = "https://calendar.app.google/HEVjuM1QFke5C7Gi6"
TURNITIN_URL = "https://www.turnitin.com/privacy"


@dataclass(frozen=True)
class Meeting:
    unit: str
    focus: str
    work: str
    prepare: str


@dataclass(frozen=True)
class Version:
    key: str
    section: str
    label: str
    days: str
    minutes: int
    meeting_time: str
    room: str
    weekdays: frozenset[int]
    meetings: tuple[Meeting, ...]
    no_class: tuple[tuple[date, str, str], ...]


def m(unit: str, focus: str, work: str, prepare: str) -> Meeting:
    return Meeting(unit, focus, work, prepare)


def time_room_text(version: Version) -> str:
    if version.meeting_time == "TBA" and version.room == "TBA":
        return "TBA · Confirm in Registrar schedule and Canvas"
    if version.meeting_time == "TBA":
        return f"Time TBA · {version.room}"
    if version.room == "TBA":
        return f"{version.meeting_time} · Room TBA"
    return f"{version.meeting_time} · {version.room}"


def schedule_assumption_text(version: Version) -> str:
    if version.meeting_time == "TBA":
        return (
            "The dates and classroom reflect the current course information. Exact clock time, final-exam "
            "appointment, assignment due times, and any later changes must be confirmed in the "
            "Registrar schedule and Canvas."
        )
    return (
        "The dates, class time, and classroom reflect the current course information. Exact final-exam "
        "appointment, assignment due times, and any later changes must be confirmed in the "
        "Registrar schedule and Canvas."
    )


MW_MEETINGS = (
    m("Orientation", "Course launch and Excel readiness", "Canvas tour, decision warm-up, workbook setup, and model habits", "Complete Start Here; activate Connect; skim Ch. 1; read Ch. 4"),
    m("TVM", "Single-sum present and future value", "Cash-flow timelines; solve PV and FV", "Complete Ch. 4 practice"),
    m("TVM", "Rates, periods, and Excel TVM functions", "Build and check a single-sum model; concept check", "Complete Ch. 4 quiz and corrections"),
    m("TVM", "Annuities and Excel", "Compare annuity types; model payment streams with PV, FV, and PMT", "Complete Ch. 5 practice I"),
    m("TVM", "Loans and credit cards", "Analyze payments, balances, and borrowing cost", "Complete Ch. 5 practice II"),
    m("TVM", "Retirement and 401(k) decisions", "Savings scenario and unit synthesis", "Submit TVM workbook checkpoint"),
    m("TVM", "TVM review", "Mixed practice, retrieval, and misconception check", "Study for the TVM exam"),
    m("Assessment", "TVM exam", "Exam", "Read Ch. 6"),
    m("Markets", "Financial markets and institutions", "Trace capital flows and interest-rate signals; exam debrief", "Complete Ch. 6 quiz; read Ch. 7"),
    m("Fixed Income", "Bond cash flows, prices, and yields", "Map bond cash flows; value coupon and zero-coupon bonds", "Complete Ch. 7 practice I"),
    m("Fixed Income", "Excel bond valuation and risk", "Build a valuation model; test rate, maturity, and credit risk", "Submit bond workbook checkpoint; complete Ch. 7 quiz"),
    m("Mortgages", "Mortgage structure", "Payment, principal, interest, and refinancing inputs", "Complete mortgage practice"),
    m("Mortgages", "Excel amortization and refinancing", "Build an amortization schedule and breakeven comparison", "Submit mortgage case"),
    m("Fixed Income", "Bond and mortgage review", "Mixed practice, retrieval, and misconception check", "Study for the bond and mortgage exam"),
    m("Assessment", "Bond and mortgage exam", "Exam", "Read Ch. 8"),
    m("Equity", "Stock valuation and the dividend discount model", "Connect required return, growth, dividends, and value", "Complete Ch. 8 practice"),
    m("Equity", "Stock valuation in Excel", "Build a scenario model and sensitivity table", "Submit stock workbook checkpoint"),
    m("Risk & Return", "Returns, dispersion, and portfolio risk", "Calculate returns and standard deviation; explore diversification", "Complete Ch. 9 practice"),
    m("Risk & Return", "Capital Asset Pricing Model", "Estimate required return using beta and market risk premium", "Complete Ch. 10 practice"),
    m("Risk & Return", "Estimating risk and return in Excel", "Analyze historical data and regression intuition", "Submit risk-and-return workbook"),
    m("Equity", "Stock and risk unit review", "Integrated case and exam practice", "Study for the stock and risk exam"),
    m("Assessment", "Stock and risk exam", "Exam and brief debrief", "Read Ch. 11"),
    m("Capital Decisions", "Cost of capital and WACC", "Estimate component costs and capital weights", "Complete Ch. 11 practice"),
    m("Capital Decisions", "NPV, IRR, and capital budgeting", "Build an Excel decision model and compare decision rules", "Submit capital-budgeting practice"),
    m("Integration", "Comprehensive final review", "Practice problems, concept map, and exam strategy", "Study for the final; complete the Canvas checklist"),
)


MWF_MEETINGS = (
    m("Orientation", "Course launch: what finance helps us decide", "Canvas tour, decision warm-up, baseline check", "Complete Start Here; activate Connect; skim Ch. 1"),
    m("Foundations", "Excel readiness and financial-statement review", "Workbook setup, formulas, references, and model habits", "Complete Excel readiness practice; read Ch. 4"),
    m("TVM", "Single-sum present and future value", "Cash-flow timelines; solve PV and FV", "Complete Ch. 4 practice I"),
    m("TVM", "Single-sum TVM in Excel", "Build and audit a PV/FV model", "Complete Excel checkpoint"),
    m("TVM", "Rates and periods", "Solve for RATE and NPER; interpret results", "Complete Ch. 4 practice II"),
    m("TVM", "Single-sum concept check", "Quiz, corrections, and transfer problem", "Read Ch. 5"),
    m("TVM", "Annuities", "Compare ordinary annuities and annuities due", "Complete Ch. 5 practice I"),
    m("TVM", "Annuities in Excel", "Model payment streams with PV, FV, and PMT", "Submit TVM workbook checkpoint"),
    m("TVM", "Loans and credit cards", "Analyze payments, balances, and borrowing cost", "Complete Ch. 5 practice II"),
    m("TVM", "Retirement and 401(k) decisions", "Savings scenario and decision memo", "Complete Ch. 5 quiz"),
    m("TVM", "TVM review", "Mixed practice, retrieval, and misconception check", "Study for the TVM exam"),
    m("Assessment", "TVM exam", "Exam", "Read Ch. 6"),
    m("Markets", "Financial markets and institutions", "Trace capital flows and interest-rate signals", "Complete Ch. 6 quiz; read Ch. 7"),
    m("Fixed Income", "Exam debrief and bond cash flows", "Correct problem areas; map coupon and principal cash flows", "Complete Ch. 7 practice I"),
    m("Fixed Income", "Bond prices", "Value coupon and zero-coupon bonds", "Complete bond price problems"),
    m("Fixed Income", "Excel bond valuation", "Build and audit a bond valuation model", "Submit bond workbook checkpoint"),
    m("Fixed Income", "Bond yields", "Calculate and interpret yield to maturity", "Complete Ch. 7 practice II"),
    m("Fixed Income", "Interest-rate risk", "Test maturity, coupon, and rate sensitivity", "Complete rate-risk questions"),
    m("Fixed Income", "Credit risk and ratings", "Compare issuer risk and rating evidence", "Complete Ch. 7 quiz I"),
    m("Mortgages", "Mortgage structure", "Identify payment, principal, interest, and term", "Complete mortgage practice I"),
    m("Mortgages", "Excel amortization", "Build and check an amortization schedule", "Submit amortization checkpoint"),
    m("Mortgages", "Refinancing decision", "Calculate breakeven and compare alternatives", "Complete mortgage case"),
    m("Fixed Income", "Bond and mortgage review", "Integrated practice and exam strategy", "Study for the bond and mortgage exam"),
    m("Assessment", "Bond and mortgage exam", "Exam", "Read Ch. 8"),
    m("Equity", "Exam debrief and stock valuation basics", "Correct problem areas; map equity cash flows", "Complete Ch. 8 practice I"),
    m("Equity", "Dividend discount model", "Connect return, growth, dividends, and value", "Complete DDM problems"),
    m("Equity", "Stock valuation in Excel", "Build a scenario model and sensitivity table", "Submit stock workbook checkpoint"),
    m("Risk & Return", "Holding-period returns", "Calculate total return across scenarios", "Complete Ch. 9 practice I"),
    m("Risk & Return", "Risk and dispersion", "Calculate variance and standard deviation", "Complete Ch. 9 practice II"),
    m("Risk & Return", "Portfolio risk", "Explore diversification and correlation", "Complete Ch. 9 quiz"),
    m("Risk & Return", "Capital Asset Pricing Model", "Estimate required return with beta and market risk premium", "Complete Ch. 10 practice I"),
    m("Risk & Return", "Estimating risk and return in Excel", "Analyze historical data and regression intuition", "Submit risk-and-return workbook"),
    m("Equity", "Integrated equity case", "Combine valuation, risk, and return evidence", "Complete Ch. 10 quiz"),
    m("Equity", "Stock and risk unit review", "Mixed practice, retrieval, and misconception check", "Study for the stock and risk exam"),
    m("Assessment", "Stock and risk exam", "Exam", "Read Ch. 11"),
    m("Capital Decisions", "Cost of capital and WACC", "Estimate component costs and capital weights", "Complete Ch. 11 practice I"),
    m("Capital Decisions", "WACC in Excel", "Build and audit a weighted-cost model", "Submit WACC checkpoint"),
    m("Capital Decisions", "NPV, IRR, and capital budgeting", "Build an Excel decision model and compare decision rules", "Complete Ch. 11 practice II"),
    m("Integration", "Comprehensive final review", "Practice problems, concept map, and exam strategy", "Study for the final; complete the Canvas checklist"),
)


VERSIONS = (
    Version(
        key="MW_75min",
        section="01",
        label="Monday/Wednesday · 2:00-3:15 p.m.",
        days="Monday and Wednesday",
        minutes=75,
        meeting_time="2:00-3:15 p.m.",
        room="GSB 155",
        weekdays=frozenset({0, 2}),
        meetings=MW_MEETINGS,
        no_class=(
            (date(2026, 9, 7), "Sep. 7", "Labor Day observed - no class"),
            (date(2026, 10, 12), "Oct. 12", "Indigenous Peoples/Columbus Day observed - no class"),
            (date(2026, 11, 23), "Nov. 23 & 25", "Thanksgiving recess - no class"),
        ),
    ),
    Version(
        key="MWF_50min",
        section="02",
        label="Monday/Wednesday/Friday · 50 minutes",
        days="Monday, Wednesday, and Friday",
        minutes=50,
        meeting_time="TBA",
        room="GSB 155",
        weekdays=frozenset({0, 2, 4}),
        meetings=MWF_MEETINGS,
        no_class=(
            (date(2026, 9, 7), "Sep. 7", "Labor Day observed - no class"),
            (date(2026, 10, 12), "Oct. 12", "Indigenous Peoples/Columbus Day observed - no class"),
            (date(2026, 11, 23), "Nov. 23-27", "Thanksgiving recess - no class"),
        ),
    ),
)


UNIT_COLORS = {
    "Orientation": (BLUE, PALE_BLUE),
    "Foundations": (BLUE, PALE_BLUE),
    "TVM": (TERRA, PALE_TERRA),
    "Assessment": (NAVY, "E2E8F0"),
    "Markets": (GOLD, PALE_GOLD),
    "Fixed Income": (TEAL, PALE_TEAL),
    "Mortgages": (TEAL, PALE_TEAL),
    "Equity": (BLUE, PALE_BLUE),
    "Risk & Return": (BLUE, PALE_BLUE),
    "Capital Decisions": (GOLD, PALE_GOLD),
    "Integration": (NAVY, "E2E8F0"),
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name: str = "Arial", size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str = LINE, size: int = 4, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    settings = edges or {edge: {} for edge in ("top", "start", "bottom", "end", "insideH", "insideV")}
    for edge, values in settings.items():
        tag = "left" if edge == "start" else "right" if edge == "end" else edge
        node = borders.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            borders.append(node)
        node.set(qn("w:val"), values.get("val", "single"))
        node.set(qn("w:sz"), str(values.get("sz", size)))
        node.set(qn("w:color"), values.get("color", color))


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_paragraph_keep(paragraph, keep_next: bool = False, keep_lines: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((r_fonts, color_node, underline))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8, color=SLATE)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    field_run._r.extend((fld_char1, instr_text, fld_char2))
    set_run_font(field_run, size=8, color=SLATE)


def style_document(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.25)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 12, 5),
        ("Heading 2", 11.5, TEAL, 8, 3),
        ("Heading 3", 10, TERRA, 6, 2),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(9)
        style.paragraph_format.space_after = Pt(2)


def configure_section(section, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)


def add_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.text = "BUS210  |  FINANCE  |  FALL 2026"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        set_run_font(run, size=7.5, color=SLATE, bold=True)
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    add_page_number(paragraph)


def add_text(document: Document, text: str, *, bold_lead: str | None = None, style=None,
             align=None, color: str | None = None, italic: bool = False) -> object:
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, color=color, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=color, italic=italic)
    return paragraph


def add_callout(document: Document, label: str, text: str, fill=PALE_GOLD, accent=GOLD) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [9840])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color=accent, size=9, start={"color": accent, "sz": 18}, top={"val": "nil"}, bottom={"val": "nil"}, end={"val": "nil"})
    set_cell_margins(cell, top=130, start=180, bottom=130, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run(f"{label}: ")
    set_run_font(lead, size=9.2, color=accent, bold=True)
    rest = paragraph.add_run(text)
    set_run_font(rest, size=9.2, color=INK)
    prevent_row_split(table.rows[0])


def add_course_header(document: Document, version: Version) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [6480, 3360])
    left, right = table.rows[0].cells
    for cell in (left, right):
        shade_cell(cell, NAVY)
        set_cell_border(cell, color=NAVY, size=0)
        set_cell_margins(cell, top=180, start=240, bottom=180, end=180)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("ENDICOTT COLLEGE  ·  GERRISH SCHOOL OF BUSINESS")
    set_run_font(run, size=8, color="BFDBFE", bold=True)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("BUS210 Finance")
    set_run_font(run, size=24, color=WHITE, bold=True)
    p = left.add_paragraph()
    run = p.add_run("Fall 2026 Course Syllabus")
    set_run_font(run, size=13, color="DBEAFE", bold=True)
    p = left.add_paragraph()
    run = p.add_run(version.label)
    set_run_font(run, size=10, color="FDE68A", bold=True)
    if PHOTO.exists():
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        shape = run.add_picture(str(PHOTO), width=Inches(2.12))
        shape._inline.docPr.set("descr", "Students collaborating on a finance analysis")
        shape._inline.docPr.set("title", "Finance analysis in action")
    prevent_row_split(table.rows[0])


def add_snapshot(document: Document, version: Version) -> None:
    document.add_heading("Course Snapshot", level=1)
    table = document.add_table(rows=4, cols=2)
    set_table_geometry(table, [4920, 4920])
    values = (
        (("Course", "BUS210 Finance · 3 credits · Lecture"), ("Section", f"BUS210-{version.section}")),
        (("Meeting pattern", f"{version.days} · {version.minutes} minutes"), ("Time / room", time_room_text(version))),
        (("Instructor", "Professor Bethany Evitts, CFA"), ("Office", "Gerrish School of Business 356")),
        (("Email", "bevitts@endicott.edu"), ("Office hours", "Use the live appointment calendar")),
    )
    for row, pair in zip(table.rows, values):
        for cell, (label, value) in zip(row.cells, pair):
            shade_cell(cell, PALE_GRAY)
            set_cell_border(cell)
            set_cell_margins(cell, top=95, start=135, bottom=95, end=135)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(label.upper())
            set_run_font(r, size=7.4, color=BLUE, bold=True)
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            if label == "Email":
                add_hyperlink(p, value, "mailto:bevitts@endicott.edu")
            elif label == "Office hours":
                add_hyperlink(p, value, OFFICE_HOURS_URL)
            else:
                r = p.add_run(value)
                set_run_font(r, size=8.8, color=INK, bold=label in {"Course", "Meeting pattern"})
        prevent_row_split(row)


def add_workflow_graphic(document: Document) -> None:
    document.add_heading("How We Learn Finance", level=1)
    add_text(document, "Each class follows a practical cycle. Exact assignments and due times appear in Canvas.")
    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [3280, 3280, 3280])
    steps = (
        ("1  PREPARE", "Read or watch, then arrive with the workbook ready.", BLUE, PALE_BLUE),
        ("2  PRACTICE", "Model the decision in Excel and explain what the result means.", TEAL, PALE_TEAL),
        ("3  PROVE", "Complete a quiz, case, checkpoint, or exam independently.", TERRA, PALE_TERRA),
    )
    for cell, (label, text, accent, fill) in zip(table.rows[0].cells, steps):
        shade_cell(cell, fill)
        set_cell_border(cell, color=accent, size=7)
        set_cell_margins(cell, top=125, start=135, bottom=125, end=135)
        p = cell.paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, size=9.2, color=accent, bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=8.2, color=INK)
    prevent_row_split(table.rows[0])


def add_learning_journey(document: Document) -> None:
    document.add_heading("Your Finance Learning Journey", level=1)
    table = document.add_table(rows=1, cols=5)
    set_table_geometry(table, [1968] * 5)
    phases = (
        ("1", "Value money", "PV · FV · loans", TERRA, PALE_TERRA),
        ("2", "Read markets", "rates · institutions", GOLD, PALE_GOLD),
        ("3", "Value debt", "bonds · mortgages", TEAL, PALE_TEAL),
        ("4", "Value equity", "stocks · risk · CAPM", BLUE, PALE_BLUE),
        ("5", "Choose projects", "WACC · NPV · IRR", NAVY, "E2E8F0"),
    )
    for cell, (number, title, detail, accent, fill) in zip(table.rows[0].cells, phases):
        shade_cell(cell, fill)
        set_cell_border(cell, color=accent, size=5)
        set_cell_margins(cell, top=95, start=90, bottom=95, end=90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(number)
        set_run_font(r, size=14, color=accent, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        set_run_font(r, size=8.2, color=INK, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(detail)
        set_run_font(r, size=6.8, color=SLATE)
    prevent_row_split(table.rows[0])


def add_academic_content(document: Document) -> None:
    document.add_heading("Course Purpose and Learning", level=1)
    document.add_heading("Catalog Description", level=2)
    add_text(document, "An introduction to finance including organization, taxes, capital markets, the commercial banking system, interest rates, financial analysis, financial forecasting, working capital management, marketable securities, accounts receivable, inventories, and short-term credit instruments. Students will use Excel to apply concepts.")
    add_text(document, "Prerequisite: ACC 175.", bold_lead="Prerequisite:")
    document.add_heading("Learning Outcomes", level=2)
    add_text(document, "The source course goal is to understand key principles in finance and, through practical problem solving, appreciate the link between theory and practice. By the end of the course, students should be able to:")
    for item in (
        "Use cash-flow timelines, financial mathematics, and Excel to solve time-value-of-money decisions.",
        "Explain how financial markets, institutions, interest rates, and risk connect savers, borrowers, and firms.",
        "Value bonds, mortgage cash flows, and common stock using appropriate assumptions and models.",
        "Measure return and risk, explain diversification, and estimate required return using CAPM.",
        "Estimate the cost of capital and evaluate investments using NPV, IRR, and sound financial judgment.",
    ):
        p = document.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=9)
    document.add_heading("Teaching and Learning Strategy", level=2)
    add_text(document, "Class combines concise instruction with active Excel-based exercises, worked examples, discussion, retrieval practice, and independent assessment. Students are expected to complete assigned preparation before class and bring the correct files to every meeting.")

    document.add_heading("Required Reading, Technology, and Files", level=1)
    for item in (
        "Finance: Applications and Theory by Cornett, Adair, and Nofsinger through McGraw Hill Connect (use the edition and access instructions posted in Canvas).",
        "Endicott Microsoft 365 access and the desktop version of Microsoft Excel installed on a PC laptop.",
        "Ability to save, organize, upload, and download Excel files without changing required filenames or formats.",
        "A financial calculator or approved calculator method when directed, plus LockDown Browser for quizzes and exams.",
    ):
        p = document.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=9)
    add_callout(document, "Bring to class", "Your charged laptop, the assigned Excel workbook, and a way to take notes. Back up important files before class.", PALE_BLUE, BLUE)


def add_evaluation(document: Document) -> None:
    document.add_heading("Evaluation and Grading", level=1)
    add_text(document, "Your course grade is based on the following components. Canvas shows the assignment-level details and current scores.")
    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [7920, 1920])
    header = table.rows[0].cells
    for cell, label in zip(header, ("Evaluation component", "Weight")):
        shade_cell(cell, NAVY)
        set_cell_border(cell, color=NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, size=8.6, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    rows = (
        ("Class participation: attendance, preparation, contribution, and professional engagement", "5%"),
        ("Readings / SmartBook activities", "5%"),
        ("Homework problems and Excel practice (completion-based unless Canvas states otherwise)", "5%"),
        ("Quizzes (lowest quiz score dropped)", "20%"),
        ("Three unit exams", "50%"),
        ("Comprehensive final exam", "15%"),
    )
    for idx, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            shade_cell(cell, WHITE if idx % 2 == 0 else PALE_GRAY)
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, size=8.5, color=INK, bold=value.endswith("%"))
            if value.endswith("%"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prevent_row_split(table.rows[-1])
    add_callout(document, "Make-up assessments", "Except in verified extreme circumstances, make-ups are not permitted. If a make-up is approved, it must be completed no later than one week after the scheduled assessment unless the professor approves a different arrangement in writing.", PALE_GOLD, GOLD)

    document.add_heading("Official Grading Scale", level=2)
    grade_table = document.add_table(rows=4, cols=6)
    set_table_geometry(grade_table, [900, 2380, 900, 2380, 900, 2380])
    grade_rows = (
        ("A", "94-100", "B-", "80-83", "D+", "67-69"),
        ("A-", "90-93", "C+", "77-79", "D", "64-66"),
        ("B+", "87-89", "C", "74-76", "D-", "60-63"),
        ("B", "84-86", "C-", "70-73", "F", "Below 60"),
    )
    for row, values in zip(grade_table.rows, grade_rows):
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            shade_cell(cell, PALE_BLUE if idx % 2 == 0 else WHITE)
            set_cell_border(cell)
            set_cell_margins(cell, top=75, start=80, bottom=75, end=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run_font(r, size=8.3, color=NAVY if idx % 2 == 0 else INK, bold=idx % 2 == 0)
        prevent_row_split(row)
    add_text(document, "WX indicates withdrawn failed when assigned under College policy.", italic=True, color=SLATE)


def add_policies(document: Document, version: Version) -> None:
    document.add_heading("Course Expectations and Policies", level=1)
    document.add_heading("Attendance and Participation", level=2)
    add_text(document, f"Students are expected to attend every scheduled {version.minutes}-minute class, arrive prepared, and participate professionally. More than one unexcused absence will reduce the attendance and participation grade. Six unexcused absences may result in a request for withdrawal. Approved accommodations and verified absences will be handled through the appropriate College process.")
    document.add_heading("Workload, Communication, and Canvas", level=2)
    add_text(document, "For each credit hour, students should plan for at least two hours of work outside class each week. For this three-credit course, that means at least six hours weekly. Check Canvas at least once per day for announcements, assignments, files, and schedule adjustments. Canvas due dates take precedence over the planning notes in the tentative schedule.")
    p = document.add_paragraph()
    p.add_run("Review the ")
    add_hyperlink(p, "Endicott undergraduate academic calendar", ACADEMIC_CALENDAR_URL)
    p.add_run(" and plan to remain available through the scheduled final-exam period.")

    document.add_heading("Academic Integrity", level=2)
    add_text(document, "Students must follow Endicott College's Academic Integrity Policy. Work submitted for assessment must be your own unless collaboration is explicitly authorized. Plagiarism, copying prior submissions, unauthorized collaboration, sharing assessment content, falsification, and other cheating are serious violations. Course violations may result in a failing course grade and will be reported through the College process.")
    add_text(document, "Proper paraphrasing means reading, thinking, interpreting, and then writing in your own words while acknowledging the source. Changing only a few words from a source is not proper paraphrasing. Ask for help before submitting work if you are unsure.")

    document.add_heading("Artificial Intelligence and LockDown Browser", level=2)
    add_text(document, "AI tools or AI-assisted software are prohibited for graded work unless Professor Evitts explicitly permits a specific use in the assignment instructions. When AI use is permitted, students must follow the stated tool, disclosure, citation, and privacy requirements. Misrepresentation of AI-generated work as one's own is academic dishonesty. LockDown Browser will be used for quizzes and exams; students are responsible for installing and testing it before an assessment.")
    p = document.add_paragraph()
    p.add_run("Read Endicott's current ")
    add_hyperlink(p, "guidance for using AI", AI_GUIDANCE_URL)
    p.add_run(". Do not upload non-public institutional information, student work, or personally identifiable information to public AI tools.")

    document.add_heading("Turnitin", level=2)
    add_text(document, "Required assignments may be submitted to Turnitin for similarity review. The tool can help identify passages that are unoriginal, incorrectly cited, or missing source information. Submissions may be retained in Turnitin's database for future similarity comparisons and other educational purposes at the instructor's discretion.")
    p = document.add_paragraph()
    add_hyperlink(p, "Review Turnitin's privacy information", TURNITIN_URL)
    p.add_run(" for additional details.")

    document.add_heading("Religious Observances", level=2)
    add_text(document, "Notify the professor within the first two weeks of the semester about specific dates on which you request relief for religious observance so that an appropriate plan can be made.")
    document.add_heading("Subject to Change", level=2)
    add_text(document, "This syllabus states the course objectives and the best estimate of what the class will cover. The professor may adjust the sequence, assignments, or schedule to respond to student learning needs, College changes, or unforeseen circumstances. Any change will be announced in Canvas.")


def add_support(document: Document) -> None:
    document.add_heading("Student Support", level=1)
    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [3280, 3280, 3280])
    cards = (
        ("Accessibility", "Students requesting accommodations should connect with the Center for Accessibility Services. General inquiries: access@endicott.edu.", ACCESSIBILITY_URL, TEAL, PALE_TEAL),
        ("Academic success", "Free content and writing tutoring, advising, and academic-support services are available through the Division of Academic Success.", ACADEMIC_SUCCESS_URL, BLUE, PALE_BLUE),
        ("Professor support", "Use office hours early. Bring the problem, workbook, or concept you want to work through.", OFFICE_HOURS_URL, GOLD, PALE_GOLD),
    )
    for cell, (title, body, url, accent, fill) in zip(table.rows[0].cells, cards):
        shade_cell(cell, fill)
        set_cell_border(cell, color=accent, size=6)
        set_cell_margins(cell, top=130, start=135, bottom=130, end=135)
        p = cell.paragraphs[0]
        r = p.add_run(title)
        set_run_font(r, size=9.2, color=accent, bold=True)
        p = cell.add_paragraph()
        r = p.add_run(body)
        set_run_font(r, size=8.1, color=INK)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_hyperlink(p, "Open current resource", url, accent)
    prevent_row_split(table.rows[0])
    document.add_heading("Pregnancy and Parental Status", level=2)
    add_text(document, "Endicott does not discriminate based on pregnancy, parental status, or related medical conditions. Qualified students may be eligible for academic adjustments that support continued participation in the course.")
    p = document.add_paragraph()
    p.add_run("For current contacts and procedures, visit ")
    add_hyperlink(p, "Endicott Title IX", TITLE_IX_URL)
    p.add_run(".")


def iter_meeting_dates(version: Version) -> list[date]:
    start = date(2026, 9, 1)
    end = date(2026, 12, 11)
    excluded = {date(2026, 9, 7), date(2026, 10, 12)}
    excluded.update(date(2026, 11, 23) + timedelta(days=i) for i in range(7))
    dates = []
    current = start
    while current <= end:
        if current.weekday() in version.weekdays and current not in excluded:
            dates.append(current)
        current += timedelta(days=1)
    if len(dates) != len(version.meetings):
        raise ValueError(f"{version.key}: {len(dates)} dates for {len(version.meetings)} meetings")
    return dates


def week_number(value: date) -> int:
    return ((value - date(2026, 8, 31)).days // 7) + 1


def date_label(value: date) -> str:
    return value.strftime("%a, %b %-d")


def schedule_rows(version: Version):
    rows = []
    for dt, meeting in zip(iter_meeting_dates(version), version.meetings):
        rows.append((dt, "meeting", str(week_number(dt)), date_label(dt), meeting))
    for dt, label, note in version.no_class:
        rows.append((dt, "no-class", str(week_number(dt)), label, note))
    return sorted(rows, key=lambda row: row[0])


def add_schedule(document: Document, version: Version) -> None:
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(section, landscape=True)
    add_header_footer(section)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Fall 2026 Tentative Schedule · {version.label}")
    set_run_font(r, size=19, color=NAVY, bold=True)
    set_paragraph_keep(p, keep_next=True)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Schedule assumptions: ")
    set_run_font(r, size=8.3, color=TERRA, bold=True)
    r = p.add_run(schedule_assumption_text(version))
    set_run_font(r, size=8.3, color=SLATE)
    table = document.add_table(rows=1, cols=6)
    # The date column is intentionally wider so long labels such as
    # "Thu, Nov 19" retain clear separation from the unit label.
    widths = [620, 1260, 1620, 2600, 3930, 3650]
    set_table_geometry(table, widths)
    headers = ("Week", "Date", "Unit", "Focus", "In-class work", "Prepare / submit")
    for cell, label in zip(table.rows[0].cells, headers):
        shade_cell(cell, NAVY)
        set_cell_border(cell, color=WHITE, size=3)
        set_cell_margins(cell, top=100, start=90, bottom=100, end=90)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if label in {"Week", "Date"} else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_run_font(r, size=7.7, color=WHITE, bold=True)
    repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for _, kind, week, label, payload in schedule_rows(version):
        row = table.add_row()
        cells = row.cells
        if kind == "no-class":
            cells[0].text = week
            cells[1].text = label
            cells[2].merge(cells[5])
            cells[2].text = payload
            for idx, cell in enumerate(cells[:3]):
                shade_cell(cell, "F1F5F9")
                set_cell_border(cell, color=LINE)
                set_cell_margins(cell, top=85, start=90, bottom=85, end=90)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=7.4, color=SLATE, bold=idx == 2, italic=True)
        else:
            meeting = payload
            values = (week, label, meeting.unit, meeting.focus, meeting.work, meeting.prepare)
            accent, fill = UNIT_COLORS[meeting.unit]
            for idx, (cell, value) in enumerate(zip(cells, values)):
                shade_cell(cell, fill if idx >= 2 else WHITE)
                set_cell_border(cell, color=LINE, size=3)
                set_cell_margins(cell, top=75, start=80, bottom=75, end=80)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in {0, 1} else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(value)
                set_run_font(r, size=7.15 if idx >= 3 else 7.25, color=accent if idx == 2 else INK, bold=idx == 2)
        prevent_row_split(row)
    row = table.add_row()
    cells = row.cells
    cells[0].text = "Finals"
    cells[1].text = "Dec. 14-18"
    cells[2].merge(cells[5])
    cells[2].text = "Comprehensive final exam · Exact date, time, and room announced by the Registrar and confirmed in Canvas"
    for idx, cell in enumerate(cells[:3]):
        shade_cell(cell, NAVY)
        set_cell_border(cell, color=WHITE, size=3)
        set_cell_margins(cell, top=100, start=90, bottom=100, end=90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=7.5, color=WHITE, bold=True)
    prevent_row_split(row)


def build_docx(version: Version, output: Path) -> None:
    document = Document()
    style_document(document)
    configure_section(document.sections[0])
    add_header_footer(document.sections[0])
    add_course_header(document, version)
    add_snapshot(document, version)
    add_workflow_graphic(document)
    add_learning_journey(document)
    add_academic_content(document)
    add_evaluation(document)
    add_policies(document, version)
    add_support(document)
    add_schedule(document, version)
    document.core_properties.title = f"BUS210 Finance Syllabus - Fall 2026 - {version.label}"
    document.core_properties.subject = "BUS210 Finance course syllabus and semester plan"
    document.core_properties.author = "Bethany Evitts"
    document.core_properties.keywords = "BUS210, Finance, Fall 2026, syllabus, Endicott College"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def ext_link(text: str, url: str) -> str:
    aria = html.escape(f"{text} (opens in a new window)", quote=True)
    return f'<a href="{html.escape(url, quote=True)}" target="_blank" rel="noopener noreferrer" aria-label="{aria}" style="color:#1d4ed8;font-weight:700;text-decoration:underline;">{html.escape(text)}</a>'


def card(title: str, body: str, accent: str, fill: str) -> str:
    return (
        f'<div style="flex:1 1 220px;padding:18px;border:1px solid #{accent};border-top:5px solid #{accent};border-radius:12px;background-color:#{fill};">'
        f'<h3 style="margin:0 0 7px;color:#{accent};font-size:18px;line-height:1.25;">{html.escape(title)}</h3>'
        f'<p style="margin:0;color:#334155;">{body}</p></div>'
    )


def html_schedule(version: Version) -> str:
    body = []
    for dt, kind, week, label, payload in schedule_rows(version):
        if kind == "no-class":
            body.append(
                f'<tr data-kind="no-class"><td style="padding:10px;border:1px solid #cbd5e1;text-align:center;">{week}</td>'
                f'<td style="padding:10px;border:1px solid #cbd5e1;text-align:center;white-space:nowrap;">{html.escape(label)}</td>'
                f'<td colspan="4" style="padding:10px;border:1px solid #cbd5e1;color:#475569;background-color:#f1f5f9;font-weight:700;text-align:center;">{html.escape(payload)}</td></tr>'
            )
            continue
        meeting = payload
        accent, fill = UNIT_COLORS[meeting.unit]
        body.append(
            f'<tr data-kind="meeting" data-date="{dt.isoformat()}">'
            f'<td style="padding:9px;border:1px solid #cbd5e1;text-align:center;">{week}</td>'
            f'<td style="padding:9px;border:1px solid #cbd5e1;text-align:center;white-space:nowrap;">{html.escape(label)}</td>'
            f'<td style="padding:9px;border:1px solid #cbd5e1;color:#{accent};background-color:#{fill};font-weight:700;">{html.escape(meeting.unit)}</td>'
            f'<td style="padding:9px;border:1px solid #cbd5e1;background-color:#{fill};">{html.escape(meeting.focus)}</td>'
            f'<td style="padding:9px;border:1px solid #cbd5e1;background-color:#{fill};">{html.escape(meeting.work)}</td>'
            f'<td style="padding:9px;border:1px solid #cbd5e1;background-color:#{fill};">{html.escape(meeting.prepare)}</td></tr>'
        )
    body.append(
        '<tr><td style="padding:10px;border:1px solid #ffffff;color:#ffffff;background-color:#0a2540;font-weight:700;text-align:center;">Finals</td>'
        '<td style="padding:10px;border:1px solid #ffffff;color:#ffffff;background-color:#0a2540;font-weight:700;text-align:center;white-space:nowrap;">Dec. 14-18</td>'
        '<td colspan="4" style="padding:10px;border:1px solid #ffffff;color:#ffffff;background-color:#0a2540;font-weight:700;">Comprehensive final exam · Exact date, time, and room announced by the Registrar and confirmed in Canvas.</td></tr>'
    )
    return (
        '<div role="region" aria-label="Fall 2026 tentative course schedule" tabindex="0" style="overflow-x:auto;margin:12px 0 30px;border:1px solid #cbd5e1;border-radius:10px;">'
        '<table style="width:100%;min-width:1080px;border-collapse:collapse;color:#172033;font-size:14px;line-height:1.35;">'
        '<caption style="padding:10px;color:#475569;background-color:#f8fafc;font-weight:700;text-align:left;">Meeting-by-meeting plan. Canvas is authoritative for assignment due dates and later changes.</caption>'
        '<thead><tr>'
        + "".join(f'<th scope="col" style="padding:10px;border:1px solid #ffffff;color:#ffffff;background-color:#0a2540;text-align:left;">{label}</th>' for label in ("Week", "Date", "Unit", "Focus", "In-class work", "Prepare / submit"))
        + '</tr></thead><tbody>' + "".join(body) + '</tbody></table></div>'
    )


def build_html(version: Version, output: Path) -> None:
    journey = "".join(
        card(title, detail, accent, fill)
        for title, detail, accent, fill in (
            ("1 · Value money", "PV, FV, annuities, loans, and retirement", TERRA, PALE_TERRA),
            ("2 · Read markets", "Institutions, interest rates, and capital flows", GOLD, PALE_GOLD),
            ("3 · Value debt", "Bonds, yields, risk, mortgages, and refinancing", TEAL, PALE_TEAL),
            ("4 · Value equity", "Stocks, risk, return, diversification, and CAPM", BLUE, PALE_BLUE),
            ("5 · Choose projects", "WACC, NPV, IRR, and capital budgeting", NAVY, "E2E8F0"),
        )
    )
    evaluation_rows = "".join(
        f'<tr><td style="padding:10px;border:1px solid #cbd5e1;">{html.escape(name)}</td><td style="padding:10px;border:1px solid #cbd5e1;font-weight:700;text-align:center;">{weight}%</td></tr>'
        for name, weight in (
            ("Class participation: attendance, preparation, contribution, and professional engagement", 5),
            ("Readings / SmartBook activities", 5),
            ("Homework problems and Excel practice", 5),
            ("Quizzes (lowest quiz score dropped)", 20),
            ("Three unit exams", 50),
            ("Comprehensive final exam", 15),
        )
    )
    learning_outcomes = "".join(
        f'<li style="margin:0 0 6px;">{html.escape(item)}</li>'
        for item in (
            "Use cash-flow timelines, financial mathematics, and Excel to solve time-value-of-money decisions.",
            "Explain how markets, institutions, interest rates, and risk connect savers, borrowers, and firms.",
            "Value bonds, mortgage cash flows, and common stock with appropriate assumptions and models.",
            "Measure return and risk, explain diversification, and estimate required return using CAPM.",
            "Estimate the cost of capital and evaluate investments using NPV, IRR, and financial judgment.",
        )
    )
    resource_list = "".join(
        f'<li style="margin:0 0 6px;">{item}</li>'
        for item in (
            "<strong>McGraw Hill Connect:</strong> <em>Finance: Applications and Theory</em> by Cornett, Adair, and Nofsinger; use the edition and access directions in Canvas.",
            "<strong>Excel:</strong> Endicott Microsoft 365 access and the desktop version of Microsoft Excel installed on a PC laptop.",
            "<strong>Files:</strong> Ability to save, organize, upload, and download Excel files without changing required filenames or formats.",
            "<strong>Assessment tools:</strong> A calculator method approved by the professor and LockDown Browser installed and tested before quizzes or exams.",
        )
    )
    grade_scale = "".join(
        f'<tr>{"".join(f"<td style=\"padding:8px;border:1px solid #cbd5e1;text-align:center;\"><strong>{html.escape(letter)}</strong> · {html.escape(score)}</td>" for letter, score in row)}</tr>'
        for row in (
            (("A", "94-100"), ("B-", "80-83"), ("D+", "67-69")),
            (("A-", "90-93"), ("C+", "77-79"), ("D", "64-66")),
            (("B+", "87-89"), ("C", "74-76"), ("D-", "60-63")),
            (("B", "84-86"), ("C-", "70-73"), ("F", "Below 60")),
        )
    )
    fragment = f'''<div style="max-width:1100px;margin:0 auto;color:#172033;font-family:Arial,Helvetica,sans-serif;font-size:16px;line-height:1.58;">
<div style="overflow:hidden;margin:0 0 26px;border-radius:18px;background-color:#0a2540;box-shadow:0 8px 24px rgba(10,37,64,0.14);">
  <div style="padding:34px 34px 30px;">
    <p style="margin:0 0 8px;color:#bfdbfe;font-size:13px;font-weight:700;letter-spacing:1.4px;text-transform:uppercase;">Endicott College · Gerrish School of Business</p>
    <h1 style="margin:0 0 8px;color:#ffffff;font-size:38px;line-height:1.12;">BUS210 Finance</h1>
    <p style="margin:0 0 6px;color:#dbeafe;font-size:22px;font-weight:700;">Fall 2026 Course Syllabus</p>
    <p style="margin:0;color:#fde68a;font-size:17px;font-weight:700;">{html.escape(version.label)}</p>
  </div>
</div>

<div style="display:flex;flex-wrap:wrap;gap:14px;margin:0 0 28px;">
  {card("Course", "BUS210 Finance · 3 credits · Lecture", BLUE, PALE_BLUE)}
  {card("Meeting pattern", html.escape(f"{version.days} · {version.minutes} minutes"), TEAL, PALE_TEAL)}
  {card("Section", html.escape(f"BUS210-{version.section}"), BLUE, PALE_BLUE)}
  {card("Time and room", html.escape(time_room_text(version)), GOLD, PALE_GOLD)}
  {card("Professor", f"Bethany Evitts, CFA · GSB 356<br>{ext_link('bevitts@endicott.edu', 'mailto:bevitts@endicott.edu')} · {ext_link('Office hours', OFFICE_HOURS_URL)}", TERRA, PALE_TERRA)}
</div>

<div style="padding:20px 22px;margin:0 0 30px;border-left:7px solid #b7791f;border-radius:10px;background-color:#fffbeb;">
  <p style="margin:0;color:#172033;"><strong style="color:#92400e;">Planning note:</strong> {html.escape(schedule_assumption_text(version))}</p>
</div>

<h2 style="margin:32px 0 10px;color:#0a2540;font-size:28px;line-height:1.2;">How we learn finance</h2>
<p style="margin:0 0 14px;color:#475569;">Each class follows a practical cycle. Exact assignments and due times appear in Canvas.</p>
<div style="display:flex;flex-wrap:wrap;gap:14px;margin:0 0 30px;">
  {card("1 · Prepare", "Read or watch, then arrive with the workbook ready.", BLUE, PALE_BLUE)}
  {card("2 · Practice", "Model the decision in Excel and explain what the result means.", TEAL, PALE_TEAL)}
  {card("3 · Prove", "Complete a quiz, case, checkpoint, or exam independently.", TERRA, PALE_TERRA)}
</div>

<h2 style="margin:32px 0 14px;color:#0a2540;font-size:28px;line-height:1.2;">Your finance learning journey</h2>
<div style="display:flex;flex-wrap:wrap;gap:12px;margin:0 0 30px;">{journey}</div>

<h2 style="margin:32px 0 10px;color:#0a2540;font-size:28px;line-height:1.2;">Course purpose and learning</h2>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Catalog description</h3>
<p style="margin:0 0 10px;">An introduction to finance including organization, taxes, capital markets, the commercial banking system, interest rates, financial analysis, financial forecasting, working capital management, marketable securities, accounts receivable, inventories, and short-term credit instruments. Students will use Excel to apply concepts.</p>
<p style="margin:0 0 14px;"><strong>Prerequisite:</strong> ACC 175.</p>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Learning outcomes</h3>
<p style="margin:0 0 8px;">The source course goal is to understand key principles in finance and, through practical problem solving, appreciate the link between theory and practice. By the end of the course, students should be able to:</p>
<ul style="margin:0 0 14px;padding-left:24px;">{learning_outcomes}</ul>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Teaching and learning strategy</h3>
<p style="margin:0 0 18px;">Class combines concise instruction with active Excel-based exercises, worked examples, discussion, retrieval practice, and independent assessment. Complete assigned preparation before class and bring the correct files to every meeting.</p>

<h2 style="margin:32px 0 10px;color:#0a2540;font-size:28px;line-height:1.2;">Required reading, technology, and files</h2>
<ul style="margin:0 0 14px;padding-left:24px;">{resource_list}</ul>
<div style="padding:18px 20px;margin:0 0 28px;border-left:7px solid #1d4ed8;border-radius:10px;background-color:#eff6ff;"><p style="margin:0;"><strong style="color:#1e40af;">Bring to class:</strong> Your charged laptop, assigned Excel workbook, and a way to take notes. Back up important files before class.</p></div>

<h2 style="margin:32px 0 10px;color:#0a2540;font-size:28px;line-height:1.2;">Evaluation and grading</h2>
<p style="margin:0 0 12px;">Canvas shows assignment-level details and current scores.</p>
<div style="overflow-x:auto;margin:0 0 18px;"><table style="width:100%;border-collapse:collapse;"><thead><tr><th scope="col" style="padding:10px;border:1px solid #ffffff;color:#ffffff;background-color:#0a2540;text-align:left;">Evaluation component</th><th scope="col" style="width:120px;padding:10px;border:1px solid #ffffff;color:#ffffff;background-color:#0a2540;text-align:center;">Weight</th></tr></thead><tbody>{evaluation_rows}</tbody></table></div>
<div style="padding:18px 20px;margin:0 0 24px;border-left:7px solid #b7791f;border-radius:10px;background-color:#fffbeb;"><p style="margin:0;"><strong style="color:#92400e;">Make-up assessments:</strong> Except in verified extreme circumstances, make-ups are not permitted. If approved, a make-up must be completed no later than one week after the scheduled assessment unless a different arrangement is approved in writing.</p></div>
<h3 style="margin:18px 0 8px;color:#0f766e;font-size:20px;">Official grading scale</h3>
<div style="overflow-x:auto;margin:0 0 8px;"><table style="width:100%;min-width:620px;border-collapse:collapse;background-color:#ffffff;"><tbody>{grade_scale}</tbody></table></div>
<p style="margin:0 0 24px;color:#475569;font-size:14px;"><em>WX indicates withdrawn failed when assigned under College policy.</em></p>

<h2 style="margin:32px 0 10px;color:#0a2540;font-size:28px;line-height:1.2;">Course expectations and policies</h2>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Attendance and participation</h3>
<p style="margin:0 0 12px;">Students are expected to attend every scheduled {version.minutes}-minute class, arrive prepared, and participate professionally. More than one unexcused absence will reduce the attendance and participation grade. Six unexcused absences may result in a request for withdrawal. Approved accommodations and verified absences will be handled through the appropriate College process.</p>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Workload, communication, and Canvas</h3>
<p style="margin:0 0 12px;">Plan for at least six hours of work outside class each week. Check Canvas at least once per day for announcements, assignments, files, and schedule adjustments. Canvas due dates take precedence over the tentative schedule. Review the {ext_link('Endicott undergraduate academic calendar', ACADEMIC_CALENDAR_URL)} and remain available through the final-exam period.</p>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Academic integrity</h3>
<p style="margin:0 0 10px;">Students must follow Endicott College's Academic Integrity Policy. Work submitted for assessment must be your own unless collaboration is explicitly authorized. Plagiarism, copying prior submissions, unauthorized collaboration, sharing assessment content, falsification, and other cheating are serious violations. Course violations may result in a failing course grade and will be reported through the College process.</p>
<p style="margin:0 0 12px;">Proper paraphrasing means reading, thinking, interpreting, and then writing in your own words while acknowledging the source. Changing only a few words is not proper paraphrasing. Ask for help before submitting work if you are unsure.</p>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Artificial intelligence and LockDown Browser</h3>
<p style="margin:0 0 12px;">AI tools or AI-assisted software are prohibited for graded work unless Professor Evitts explicitly permits a specific use in the assignment instructions. When permitted, follow the stated tool, disclosure, citation, and privacy requirements. Misrepresenting AI-generated work as your own is academic dishonesty. LockDown Browser will be used for quizzes and exams; install and test it before an assessment. Read Endicott's {ext_link('current AI guidance', AI_GUIDANCE_URL)}. Do not upload non-public institutional information, student work, or personally identifiable information to public AI tools.</p>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Turnitin</h3>
<p style="margin:0 0 12px;">Required assignments may be submitted to Turnitin for similarity review. Submissions may be retained in Turnitin's database for future comparisons and other educational purposes at the instructor's discretion. {ext_link("Review Turnitin's privacy information", TURNITIN_URL)}.</p>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Religious observances</h3>
<p style="margin:0 0 12px;">Notify the professor within the first two weeks about specific dates on which you request relief for religious observance so an appropriate plan can be made.</p>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Subject to change</h3>
<p style="margin:0 0 24px;">This syllabus states the course objectives and the best estimate of what the class will cover. The professor may adjust the sequence, assignments, or schedule in response to learning needs, College changes, or unforeseen circumstances. Changes will be announced in Canvas.</p>

<h2 style="margin:32px 0 12px;color:#0a2540;font-size:28px;line-height:1.2;">Student support</h2>
<div style="display:flex;flex-wrap:wrap;gap:14px;margin:0 0 18px;">
  {card("Accessibility", f"Request accommodations through the Center for Accessibility Services. General inquiries: access@endicott.edu.<br>{ext_link('Open Accessibility Services', ACCESSIBILITY_URL)}", TEAL, PALE_TEAL)}
  {card("Academic success", f"Free content and writing tutoring, advising, and other academic-support services are available.<br>{ext_link('Open the Division of Academic Success', ACADEMIC_SUCCESS_URL)}", BLUE, PALE_BLUE)}
  {card("Professor support", f"Use office hours early. Bring the problem, workbook, or concept you want to work through.<br>{ext_link('Book office hours', OFFICE_HOURS_URL)}", GOLD, PALE_GOLD)}
</div>
<h3 style="margin:18px 0 6px;color:#0f766e;font-size:20px;">Pregnancy and parental status</h3>
<p style="margin:0 0 28px;">Endicott does not discriminate based on pregnancy, parental status, or related medical conditions. Qualified students may be eligible for academic adjustments. For current contacts and procedures, visit {ext_link('Endicott Title IX', TITLE_IX_URL)}.</p>

<h2 style="margin:34px 0 8px;color:#0a2540;font-size:28px;line-height:1.2;">Fall 2026 tentative schedule</h2>
<p style="margin:0 0 12px;color:#475569;">{html.escape(version.label)} · Classes begin September 1; Thanksgiving recess follows November 20 and classes resume November 30; last class day is December 11; final exams are December 14-18.</p>
{html_schedule(version)}
</div>'''
    output.parent.mkdir(parents=True, exist_ok=True)
    # Canvas fragments intentionally omit document-level metadata. Convert
    # non-ASCII punctuation to numeric entities so local previews and Canvas
    # both render the same way even when the host page controls the charset.
    output.write_text(fragment.encode("ascii", "xmlcharrefreplace").decode("ascii"), encoding="ascii")


def validate_html(path: Path, version: Version) -> None:
    text = path.read_text(encoding="utf-8")
    lowered = text.lower()
    for forbidden in ("<!doctype", "<html", "<head", "<body", "<style", "<script", "<link"):
        if forbidden in lowered:
            raise ValueError(f"{path.name}: forbidden Canvas wrapper/tag {forbidden}")
    if len(re.findall(r"<h1\b", lowered)) != 1:
        raise ValueError(f"{path.name}: expected exactly one h1")
    meeting_count = len(re.findall(r'data-kind="meeting"', text))
    if meeting_count != len(version.meetings):
        raise ValueError(f"{path.name}: expected {len(version.meetings)} meeting rows, found {meeting_count}")
    for match in re.finditer(r'<a\s+([^>]+)>', text, re.I):
        attrs = match.group(1).lower()
        if 'href="http' in attrs and not all(token in attrs for token in ('target="_blank"', 'noopener', 'noreferrer', 'aria-label=')):
            raise ValueError(f"{path.name}: external link missing safe new-window attributes")
    if sum((5, 5, 5, 20, 50, 15)) != 100:
        raise ValueError("grading weights do not total 100")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=OUTPUT_DIR)
    args = parser.parse_args()
    for version in VERSIONS:
        docx_path = args.output_dir / f"BUS210_{version.section}_Finance_Syllabus_FA2026_{version.key}.docx"
        html_path = args.output_dir / f"BUS210_Finance_Syllabus_FA2026_{version.key}_Canvas.html"
        build_docx(version, docx_path)
        build_html(version, html_path)
        validate_html(html_path, version)
        dates = iter_meeting_dates(version)
        print(f"built {docx_path.name}: {len(dates)} meetings, {dates[0]}, {dates[-1]}")
        print(f"built {html_path.name}: Canvas validation passed")


if __name__ == "__main__":
    main()
